"""
convert_tools.py - Document conversion endpoints (lintas format umum).
Untuk operasi PDF spesifik yang lebih detail, lihat pdf_tools.py.
Untuk operasi spreadsheet detail, lihat spreadsheet_tools.py.
"""
import json
import os
import zipfile

from flask import Blueprint, request

from utils import lazy_import, save_upload, new_tmp_dir, error_response, send_file_response, ok_response

convert_bp = Blueprint("convert_tools", __name__, url_prefix="/api/convert")


@convert_bp.route("/pdf-to-docx", methods=["POST"])
def pdf_to_docx():
    """PDF -> Word (.docx), termasuk tabel. Form field: file."""
    f = request.files.get("file")
    if not f:
        return error_response("Field 'file' wajib diisi.")

    pdfplumber = lazy_import("pdfplumber")
    lazy_import("docx", "python-docx")
    from docx import Document

    tmp = new_tmp_dir()
    path = save_upload(f, tmp)
    doc = Document()

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            doc.add_heading(f"Halaman {i}", level=2)
            for para in (page.extract_text() or "").split("\n"):
                if para.strip():
                    doc.add_paragraph(para)
            for t_idx, table in enumerate(page.extract_tables(), start=1):
                if not table:
                    continue
                rows, cols = len(table), max(len(r) for r in table)
                wt = doc.add_table(rows=rows, cols=cols)
                wt.style = "Table Grid"
                for r, row in enumerate(table):
                    for c, cell in enumerate(row):
                        wt.cell(r, c).text = str(cell) if cell else ""

    out_path = os.path.join(tmp, "converted.docx")
    doc.save(out_path)
    return send_file_response(out_path, "converted.docx")


@convert_bp.route("/pdf-to-txt", methods=["POST"])
def pdf_to_txt():
    """PDF -> teks polos. Form field: file."""
    f = request.files.get("file")
    if not f:
        return error_response("Field 'file' wajib diisi.")

    pdfplumber = lazy_import("pdfplumber")
    tmp = new_tmp_dir()
    path = save_upload(f, tmp)

    chunks = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            chunks.append(f"--- Halaman {i} ---\n{page.extract_text() or ''}\n")

    out_path = os.path.join(tmp, "converted.txt")
    with open(out_path, "w", encoding="utf-8") as fo:
        fo.write("\n".join(chunks))

    return send_file_response(out_path, "converted.txt", "text/plain")


@convert_bp.route("/pdf-to-images", methods=["POST"])
def pdf_to_images():
    """PDF -> gambar per halaman (zip). Form field: file, format (jpg/png), dpi."""
    f = request.files.get("file")
    if not f:
        return error_response("Field 'file' wajib diisi.")

    fmt = request.form.get("format", "jpg").lower().replace("jpeg", "jpg")
    dpi = int(request.form.get("dpi", 200))
    pil_fmt = "JPEG" if fmt == "jpg" else fmt.upper()

    pdf2image = lazy_import("pdf2image")
    from pdf2image import convert_from_path

    tmp = new_tmp_dir()
    path = save_upload(f, tmp)
    images = convert_from_path(path, dpi=dpi)

    zip_path = os.path.join(tmp, "pages.zip")
    with zipfile.ZipFile(zip_path, "w") as zf:
        for i, img in enumerate(images, start=1):
            if pil_fmt == "JPEG" and img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            page_path = os.path.join(tmp, f"page_{i:03d}.{fmt}")
            img.save(page_path, pil_fmt)
            zf.write(page_path, arcname=os.path.basename(page_path))

    return send_file_response(zip_path, "pages.zip", "application/zip")


@convert_bp.route("/images-to-pdf", methods=["POST"])
def images_to_pdf():
    """Banyak gambar -> 1 PDF. Form field: files (multiple, urut sesuai upload)."""
    img2pdf = lazy_import("img2pdf")
    files = request.files.getlist("files")
    if not files:
        return error_response("Minimal 1 file gambar (field 'files').")

    tmp = new_tmp_dir()
    paths = [save_upload(f, tmp) for f in files]

    out_path = os.path.join(tmp, "converted.pdf")
    with open(out_path, "wb") as fo:
        fo.write(img2pdf.convert(paths))

    return send_file_response(out_path, "converted.pdf", "application/pdf")


@convert_bp.route("/markdown-to-html", methods=["POST"])
def markdown_to_html():
    """Markdown -> HTML. Body: {markdown_text}."""
    markdown_lib = lazy_import("markdown")
    body = request.get_json(silent=True) or {}
    text = body.get("markdown_text", "")
    html = markdown_lib.markdown(text, extensions=["tables", "fenced_code", "codehilite"])
    return ok_response({"result": html})


@convert_bp.route("/yaml-to-json", methods=["POST"])
def yaml_to_json():
    """YAML -> JSON. Body: {yaml_text}."""
    yaml = lazy_import("yaml", "pyyaml")
    body = request.get_json(silent=True) or {}
    text = body.get("yaml_text", "")
    try:
        data = yaml.safe_load(text)
        return ok_response({"result": json.dumps(data, indent=2, ensure_ascii=False)})
    except yaml.YAMLError as e:
        return error_response(f"YAML tidak valid: {e}")


@convert_bp.route("/json-to-yaml", methods=["POST"])
def json_to_yaml():
    """JSON -> YAML. Body: {json_text}."""
    yaml = lazy_import("yaml", "pyyaml")
    body = request.get_json(silent=True) or {}
    text = body.get("json_text", "")
    try:
        data = json.loads(text)
        return ok_response({"result": yaml.dump(data, allow_unicode=True, sort_keys=False)})
    except json.JSONDecodeError as e:
        return error_response(f"JSON tidak valid: {e}")


@convert_bp.route("/csv-to-json", methods=["POST"])
def csv_to_json():
    """File CSV -> JSON. Form field: file."""
    pd = lazy_import("pandas")
    f = request.files.get("file")
    if not f:
        return error_response("Field 'file' wajib diisi.")

    tmp = new_tmp_dir()
    path = save_upload(f, tmp)
    df = pd.read_csv(path)
    return ok_response(json.loads(df.to_json(orient="records")))


@convert_bp.route("/json-to-csv", methods=["POST"])
def json_to_csv():
    """JSON array of objects (file upload) -> CSV. Form field: file."""
    pd = lazy_import("pandas")
    f = request.files.get("file")
    if not f:
        return error_response("Field 'file' wajib diisi.")

    tmp = new_tmp_dir()
    path = save_upload(f, tmp)
    with open(path, "r", encoding="utf-8") as fo:
        data = json.load(fo)

    df = pd.DataFrame(data)
    out_path = os.path.join(tmp, "converted.csv")
    df.to_csv(out_path, index=False)
    return send_file_response(out_path, "converted.csv", "text/csv")
